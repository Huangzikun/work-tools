"""
根据成绩单中的成绩，设置签名并保存成docx文件
默认的学号和分数列名和学习通保持一致

处理逻辑：
1. 读取成绩单数据，获得学生学号和分数两列
2. 创建临时目录
3. 将doc文件转换成docx文件，并存放于临时目录中
4. 根据学号查找临时目录中对应的docx文件，打开文件并签名，将文件保存至指定目录
5. 生成未找到学生的txt文件
6. 删除临时目录
"""
import os
from docx import Document
from docx.shared import Cm
from pathlib import Path
from win32com import client as wc
import platform
import argparse
import pandas as pd

# 存放第一次转换的临时文件夹名称
temp_dir_name = "temp"


# 获取成绩单的学生学号和分数两列的dataframe
def get_score_from_file(file_path, sid_col, score_col,skip_rows=0):
    df = pd.read_excel(file_path, skiprows=skip_rows,dtype=str)

    # 检查列序号是否正确
    sid_col = int(sid_col) - 1
    score_col = int(score_col) - 1   # 列序号从0开始计数
    if sid_col >= len(df.columns) or score_col >= len(df.columns):
        print("列序号错误，请检查输入参数")
        return None

    # 取出学号和分数两列
    data = pd.DataFrame(df.iloc[:, [sid_col, score_col]])
    data.columns = ['sid','score']
    data.reset_index(drop=True, inplace=True)  # 重置索引

    return data

# 获取指定目录下所有 doc/docx 文件
def get_file_list(directory_path, current_deep=0, deep=5):
    if current_deep >= deep:
        return []

    real_file_list = []
    file_list = os.listdir(directory_path)
    for file in file_list:
        if os.path.isdir(os.path.join(directory_path, file)):
            return get_file_list(os.path.join(directory_path, file), current_deep + 1, deep)
        if file.endswith(".doc") or file.endswith(".docx"):
            real_file_list.append(os.path.join(directory_path, file))

    return real_file_list


def doc_to_docx(file_path,save_path):

    # doc = word.Documents.Open(file_path)
    word = wc.Dispatch("kwps.Application")  # WPS的COM接口
    word.Visible = False
    try:
        doc = word.Documents.Open(str(Path(file_path).resolve()))
    except Exception as e:
        print(f"文件[{file_path}]打开失败. 错误信息: {e} \n")
        exit(0)
    filename = os.path.split(file_path)[1]

    save_file_path = save_path + os.sep + filename.split('.')[0] + ".docx"

    print(file_path + " save to " + save_file_path)
    doc.SaveAs(save_file_path, 12, False, "", True, "", False, False, False, False)
    doc.Close()


def sign_by_picture(file_path,save_path,score):
    doc = Document(file_path)
    if len(doc.tables) <= 0:
        print(f"can't sign file. file = {file_path}")
        return

    for table in doc.tables:
        for row in table.rows:
            for i in range(len(row.cells)):

                if str.strip(row.cells[i].text) == '教师评阅':
                    level = ""
                    score = int(score)
                    if score >= 90:
                        level = "优秀"
                    elif score >= 80:
                        level = "良好"
                    elif score >= 60:
                        level = "及格"
                    else:
                        level = "不及格！"

                    row.cells[i + 1].text = (f"{level} 成绩：{score}分"
                                             f"{os.linesep}"
                                             f"{os.linesep}"
                                             f"{os.linesep}"
                                             f"{os.linesep}"
                                             f"{os.linesep}"
                                             f"                       教师签名：{sign}   {sign_date}")
                    row.cells[i + 1].add_paragraph().add_run().add_picture(sign_picture, width=Cm(2))
                    doc.save(save_path)
                    print(f"sign successfully. file = {file_path}")

                    return

    print(f"sign fail. file = {file_path}")


'''
执行内容
'''
parser = argparse.ArgumentParser()
parser.add_argument('--applicatioon', type=str, help='原始文件目录', default="kwps.Application")
parser.add_argument('--base_path', type=str, help='原始文件目录', required=True)
parser.add_argument('--save_path', type=str, help='保存文件目录', required=True)
parser.add_argument('--sign_picture', type=str, help='签名图片地址', required=True)
parser.add_argument('--sign', type=str, help='签名字符串，如某某某', required=True)
parser.add_argument('--sign_date', type=str, help='签名时间', required=True)
parser.add_argument('--transcript_path', type=str, help='请输入成绩单文件路径', default=False)
parser.add_argument('--sid_col', type=str, help='学生学号列号，从1开始计数', default=1)
parser.add_argument('--score_col', type=str, help='分数列名，从1开始计数', default=9)
parser.add_argument('--data_row', type=str, help='数据所在行', default=3)

#获取参数
args = parser.parse_args()
applicatioon = args.applicatioon
base_path = args.base_path
save_path = args.save_path
sign_picture = args.sign_picture
sign = args.sign
sign_date = args.sign_date
transcript_path = args.transcript_path
sid_col = args.sid_col
score_col = args.score_col
data_row = args.data_row

os_platform = platform.system()
error_message = ""
print(f"当前运行平台: {os_platform}")

# 打开word
if applicatioon=="kwps.Application" or applicatioon=="Word.Application":
    word = wc.Dispatch(applicatioon)
else:
    print("请确认输入的--applicatioon参数是否正确 默认为WPS，Office Word请输入Word.Application")
    exit(0)

if os_platform != 'Windows':
    print("格式转换必须在Windows平台上进行且需要包含Office套件以支持doc文件. 如果只需要签名请添加 --only_sign=True")
    exit(0)

print("1.开始获取成绩单数据")
if transcript_path:
    score_df = get_score_from_file(transcript_path, sid_col, score_col,skip_rows=data_row-1)
else:
    score_df = None

if score_df is None:
    print("成绩单数据获取失败，请检查输入参数")
    exit(0)

print("2.开始创建临时目录")
temp_path = os.path.join(save_path, temp_dir_name)
if not os.path.exists(temp_path):
    os.mkdir(temp_path)

print("3.开始转换doc->docx")
print("未加工的文件所在路径为: "+base_path)
print("转换后的文件保存路径为: "+temp_path)
student_file_list = os.listdir(base_path)
for file in student_file_list:
    try:
        doc_to_docx(os.path.join(base_path, file), temp_path)
    except Exception as e:
        print(f"文件{file}转换失败. 错误信息: {e} \n")
        error_message += f"文件{file}转换失败. 错误信息: {e} \n"

print("开始转换docx转换完成.")
print("4.开始签署")

student_file_list = os.listdir(temp_path)
count = 0
# 遍历学生成绩单数据score_df，根据学号一列查找对应的文件，打开文件并前面
# 只要文件名包含对应学号，则进行签署，文件名可以包含其他内容
for row in score_df.itertuples():
    student_id = str(getattr(row, "sid"))
    find_flag = False
    # 依次检查文件名是否包含学号
    for file in student_file_list:
        if student_id != "" and student_id in file:
            find_flag = True
            file_path = os.path.join(temp_path, file)  # 保存文件名
            count += 1
            try:
                sign_by_picture(file_path,os.path.join(save_path, file), getattr(row, "score"))
            except Exception as e:
                error_message += f"文件{file_path}签署失败. 错误信息: {e} \n"
                print(f"文件{file_path}签署失败. 错误信息: {e} \n")

    if not find_flag:
        error_message += f"未找到学号为{student_id}的学生文件.  \n"

print("5.开始删除临时目录,清空其中文件")
for file in os.listdir(temp_path):
    os.remove(os.path.join(temp_path, file))
os.rmdir(temp_path)

print(f"共签署{count}个学生文件.")
error_message += f"共签署{count}个学生文件. \n"
word.Quit()

if error_message:
    print(f"签署失败. 错误信息: {error_message}")
    # 生成未找到学生的txt文件
    error_file_path = os.path.join(save_path, "error.txt")
    with open(error_file_path, "w", encoding="utf-8") as f:
        f.write(error_message)
else:
    print("签署成功.")





