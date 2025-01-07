import os
from docx import Document
from docx.shared import Cm
import platform

# 根据操作系统类型引入不同的包
if platform.system() == 'Windows':
    from win32com import client as wc

import platform
import argparse


def get_file_list(directory_path, current_deep=0, deep=5):
    if current_deep >= deep:
        return []

    real_file_list = []
    file_list = os.listdir(directory_path)
    for file in file_list:
        if os.path.isdir(os.path.join(directory_path, file)):
            return get_file_list(os.path.join(directory_path, file), current_deep + 1, deep)
        if file.endswith(".doc") or file.endswith(".docx"):
            real_file_list.append(os.path.join(directory_path, file))

    return real_file_list


def doc_to_docx(file_path):
    word = wc.Dispatch("Word.Application")
    doc = word.Documents.Open(file_path)
    filename = os.path.split(file_path)[1]
    student_path = os.path.split(file_path)[0].replace(base_path, '').split(os.sep)[1]

    student_save_path = os.path.join(save_path, student_path)

    if not os.path.exists(student_save_path):
        os.mkdir(student_save_path)

    save_file_path = student_save_path + os.sep + filename.split('.')[0] + ".docx"

    print(file_path + " save to " + save_file_path)
    doc.SaveAs(save_file_path, 12, False, "", True, "", False, False, False, False)
    doc.Close()
    word.Quit()


def sign_by_picture(file_path):
    doc = Document(file_path)
    if len(doc.tables) <= 0:
        print(f"can't sign file. file = {file_path}")
        return

    for table in doc.tables:
        for row in table.rows:
            for i in range(len(row.cells)):

                if str.strip(row.cells[i].text) == '教师评阅':
                    row.cells[i + 1].text = (f"已阅"
                                             f"{os.linesep}"
                                             f"{os.linesep}"
                                             f"{os.linesep}"
                                             f"{os.linesep}"
                                             f"{os.linesep}"
                                             f"                       教师签名：{sign}   {sign_date}")
                    row.cells[i + 1].add_paragraph().add_run().add_picture(sign_picture, width=Cm(2))
                    doc.save(file_path)
                    print(f"sign successfully. file = {file_path}")

                    return

    print(f"sign fail. file = {file_path}")


'''
执行内容
'''
parser = argparse.ArgumentParser()
parser.add_argument('--base_path', type=str, help='原始文件目录', required=True)
parser.add_argument('--sign_picture', type=str, help='签名图片地址', required=True)
parser.add_argument('--sign', type=str, help='签名字符串，如某某某', required=True)
parser.add_argument('--sign_date', type=str, help='签名时间', required=True)
args = parser.parse_args()

#获取参数
base_path = args.base_path
sign_picture = args.sign_picture
sign = args.sign
sign_date = args.sign_date


os_platform = platform.system()
print(f"当前运行平台: {os_platform}")
print("开始签署")

student_file_list = os.listdir(base_path)

error_file_list = []
for student in student_file_list:

    if student.startswith("."):
        continue
    file_list = get_file_list(os.path.join(base_path, student))
    for file in file_list:
        if file.endswith(".doc"):
            error_file_list.append(os.path.join(base_path, student, file))
            continue
        sign_by_picture(file)

print("签署完成.")
print(error_file_list)


